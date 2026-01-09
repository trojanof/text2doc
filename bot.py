import logging
import tempfile
import os
import telebot
from docx import Document

# Настройка логирования
logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)
logger = logging.getLogger(__name__)

# Токен бота (замените на ваш токен)
BOT_TOKEN = os.environ['TOKEN']

# Создаем бота
bot = telebot.TeleBot(BOT_TOKEN)


@bot.message_handler(commands=['start'])
def send_welcome(message):
    """Обработчик команды /start"""
    welcome_text = (
        "Привет! 👋\n\n"
        "Отправь мне сообщение с форматированным текстом "
        "(жирный, курсив, подчеркнутый, зачеркнутый), "
        "и я сохраню его в DOCX файл!\n\n"
        "Поддерживаемые форматы:\n"
        "• **жирный текст**\n"
        "• *курсив*\n"
        "• __подчеркнутый__\n"
        "• ~~зачеркнутый~~\n"
        "• `моноширинный`\n"
        "• ```блок кода```"
    )
    bot.reply_to(message, welcome_text)


def apply_formatting(run, entity_type):
    """Применяет форматирование к тексту в зависимости от типа"""
    if entity_type == "bold":
        run.bold = True
    elif entity_type == "italic":
        run.italic = True
    elif entity_type == "underline":
        run.underline = True
    elif entity_type == "strikethrough":
        run.font.strike = True
    elif entity_type == "code":
        run.font.name = "Courier New"
    elif entity_type == "pre":
        run.font.name = "Courier New"


def create_docx_with_formatting(text, entities):
    """Создает DOCX документ с форматированием"""
    doc = Document()
    paragraph = doc.add_paragraph()

    # Если нет entities, просто добавляем обычный текст
    if not entities:
        paragraph.add_run(text)
        return doc

    # Сортируем entities по offset
    sorted_entities = sorted(entities, key=lambda x: x.offset)

    current_pos = 0

    for entity in sorted_entities:
        # Добавляем обычный текст до форматированного
        if entity.offset > current_pos:
            paragraph.add_run(text[current_pos:entity.offset])

        # Добавляем форматированный текст
        formatted_text = text[entity.offset:entity.offset + entity.length]
        run = paragraph.add_run(formatted_text)
        apply_formatting(run, entity.type)

        current_pos = entity.offset + entity.length

    # Добавляем оставшийся текст
    if current_pos < len(text):
        paragraph.add_run(text[current_pos:])

    return doc


@bot.message_handler(content_types=['text'])
def handle_text_message(message):
    """Обработчик текстовых сообщений"""
    try:
        text = message.text
        entities = message.entities or []

        if not text:
            bot.reply_to(message, "Отправьте текстовое сообщение!")
            return

        # Отправляем уведомление о начале обработки
        processing_msg = bot.reply_to(message, "Создаю DOCX файл... ⏳")

        # Создаем DOCX документ
        doc = create_docx_with_formatting(text, entities)

        # Сохраняем во временный файл
        with tempfile.NamedTemporaryFile(suffix='.docx',
                                         delete=False) as temp_file:
            doc.save(temp_file.name)
            temp_filename = temp_file.name

        try:
            # Отправляем документ
            with open(temp_filename, 'rb') as doc_file:
                bot.send_document(
                    message.chat.id,
                    doc_file,
                    caption="Ваш форматированный текст сохранен в DOCX! 📄",
                    reply_to_message_id=message.message_id
                )

            # Удаляем сообщение о процессе обработки
            bot.delete_message(message.chat.id, processing_msg.message_id)

        finally:
            # Удаляем временный файл
            os.unlink(temp_filename)

    except Exception as e:
        logger.error(f"Ошибка при обработке сообщения: {e}")
        bot.reply_to(
            message,
            "Произошла ошибка при создании документа. "
            "Попробуйте еще раз или обратитесь к администратору."
        )


@bot.message_handler(content_types=['document'])
def handle_document(message):
    """Обработчик документов"""
    bot.reply_to(
        message,
        "Я работаю только с текстовыми сообщениями! "
        "Отправьте мне текст с форматированием."
    )


@bot.message_handler(content_types=['photo', 'video', 'audio', 'voice',
                                    'sticker'])
def handle_media(message):
    """Обработчик медиа-файлов"""
    bot.reply_to(
        message,
        "Я обрабатываю только текстовые сообщения! "
        "Отправьте мне текст с форматированием."
    )


def main():
    """Основная функция"""
    logger.info("Бот запущен...")
    logger.info(f"Бот: @{bot.get_me().username}")
    logger.info("Бот работает по адресу https://t.me/msg2doc_bot")

    # Запускаем бота
    try:
        bot.polling(none_stop=True, interval=0, timeout=20)
    except Exception as e:
        logger.error(f"Ошибка при запуске бота: {e}")
        print("Перезапуск бота через 5 секунд...")
        import time
        time.sleep(5)
        main()


if __name__ == '__main__':
    main()
